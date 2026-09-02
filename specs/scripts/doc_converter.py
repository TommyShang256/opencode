"""DIS集成需求文档转Word工具

将Markdown格式的DIS集成需求文档转换为美观的Word文档，
确保标题层级正确（左侧导航栏可见）。

使用方法:
    python doc_converter.py <输入md文件> [输出docx文件]
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import NamedTuple

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: Please install python-docx")
    print("Run: pip install python-docx")
    sys.exit(1)


class StyleConfig(NamedTuple):
    """样式配置"""
    h1_size: int = 22
    h1_color: tuple = (0, 51, 102)  # 深蓝
    h2_size: int = 18
    h2_color: tuple = (0, 82, 163)  # 中蓝
    h3_size: int = 16
    h3_color: tuple = (51, 102, 204)  # 浅蓝
    h4_size: int = 14
    h4_color: tuple = (0, 0, 0)
    h5_size: int = 13
    h5_color: tuple = (64, 64, 64)
    h6_size: int = 12
    h6_color: tuple = (96, 96, 96)
    body_font: str = "微软雅黑"
    body_size: int = 11
    code_font: str = "Consolas"
    code_size: int = 10
    table_header_bg: tuple = (230, 242, 255)
    table_header_color: tuple = (0, 51, 102)


class DisDocConverter:
    """DIS文档转换器 - 使用Word内置标题样式"""

    def __init__(self, style_config: StyleConfig | None = None):
        self.config = style_config or StyleConfig()
        self.doc = Document()
        self._table_rows: list[list[str]] = []
        self._in_table = False

        self._setup_page()
        self._setup_heading_styles()

    def _rgb(self, color: tuple) -> RGBColor:
        return RGBColor(*color)

    def _setup_page(self) -> None:
        """设置页面布局"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _setup_heading_styles(self) -> None:
        """设置Word内置标题样式"""
        sizes = {
            'Heading 1': (self.config.h1_size, self.config.h1_color),
            'Heading 2': (self.config.h2_size, self.config.h2_color),
            'Heading 3': (self.config.h3_size, self.config.h3_color),
            'Heading 4': (self.config.h4_size, self.config.h4_color),
            'Heading 5': (self.config.h5_size, self.config.h5_color),
            'Heading 6': (self.config.h6_size, self.config.h6_color),
        }

        for style_name, (size, color) in sizes.items():
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                style.font.name = self.config.body_font
                style.font.size = Pt(size)
                style.font.bold = True
                style.font.color.rgb = self._rgb(color)

                # 确保中文字体生效
                rPr = style.element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), self.config.body_font)
                rPr.append(rFonts)

        # 设置正文样式
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = self.config.body_font
        normal_style.font.size = Pt(self.config.body_size)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal_style.paragraph_format.space_after = Pt(6)

        rPr = normal_style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), self.config.body_font)
        rPr.append(rFonts)

    def _add_heading(self, text: str, level: int) -> None:
        """添加标题 - 使用Word内置标题样式"""
        style_name = f'Heading {level}'

        if style_name in self.doc.styles:
            para = self.doc.add_paragraph(text, style=style_name)
        else:
            # 回退：手动设置样式
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = self.config.body_font
            run.font.bold = True

            sizes = {1: 22, 2: 18, 3: 16, 4: 14, 5: 13, 6: 12}
            colors = {
                1: (0, 51, 102), 2: (0, 82, 163), 3: (51, 102, 204),
                4: (0, 0, 0), 5: (64, 64, 64), 6: (96, 96, 96)
            }
            run.font.size = Pt(sizes.get(level, 12))
            run.font.color.rgb = self._rgb(colors.get(level, (0, 0, 0)))

        # 一二级标题添加下划线
        if level <= 2:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '0070C0')
            pBdr.append(bottom)
            pPr.append(pBdr)

    def _add_paragraph(self, text: str) -> None:
        """添加普通段落"""
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        self._parse_rich_text(para, text)

    def _parse_rich_text(self, para, text: str) -> None:
        """解析富文本"""
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)

        for part in parts:
            run = para.add_run()

            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run.text = part[1:-1]
                run.font.name = self.config.code_font
                run.font.color.rgb = self._rgb((128, 0, 128))
            else:
                run.text = part
                run.font.name = self.config.body_font

            run.font.size = Pt(self.config.body_size)

    def _add_list_item(self, text: str) -> None:
        """添加列表项"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.first_line_indent = Cm(-0.3)

        run = para.add_run("• " + text)
        run.font.name = self.config.body_font
        run.font.size = Pt(self.config.body_size)

    def _flush_table(self) -> None:
        """输出表格"""
        if not self._table_rows:
            return

        rows = len(self._table_rows)
        cols = max(len(row) for row in self._table_rows) if self._table_rows else 0

        if rows > 0 and cols > 0:
            table = self.doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_data in enumerate(self._table_rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text

                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.name = self.config.body_font
                                r.font.size = Pt(self.config.code_size if col_idx > 0 else self.config.body_size)
                                if row_idx == 0:
                                    r.bold = True
                                    r.font.color.rgb = self._rgb(self.config.table_header_color)

                            if row_idx == 0:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)

                        # 表头背景色
                        if row_idx == 0:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), '%02X%02X%02X' % self.config.table_header_bg)
                            tcPr.append(shd)

            self.doc.add_paragraph()

        self._table_rows = []
        self._in_table = False

    def _add_separator(self) -> None:
        """添加分隔线"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _process_line(self, line: str) -> None:
        """处理一行Markdown内容"""
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            self._flush_table()
            return

        # 表格处理
        if stripped.startswith('|'):
            self._in_table = True
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            self._table_rows.append(cells)
            return

        # 代码块
        if stripped.startswith('```'):
            return

        # 标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            if 1 <= level <= 6:
                text = stripped.lstrip('#').strip()
                self._flush_table()
                self._add_heading(text, level)
                return

        # 分隔线
        if stripped.startswith('---') or stripped.startswith('***'):
            self._flush_table()
            self._add_separator()
            return

        # 列表项
        if stripped.startswith('- ') or stripped.startswith('* '):
            self._flush_table()
            self._add_list_item(stripped[2:])
            return

        # 普通段落
        self._flush_table()
        self._add_paragraph(stripped)

    def convert(self, markdown: str) -> Document:
        """转换Markdown内容"""
        lines = markdown.split('\n')

        for line in lines:
            self._process_line(line)

        if self._in_table:
            self._flush_table()

        return self.doc

    def convert_file(self, input_path: str, output_path: str | None = None) -> str:
        """转换文件"""
        input_path = Path(input_path)

        if not input_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")

        content = input_path.read_text(encoding='utf-8')
        self.convert(content)

        if output_path is None:
            output_path = input_path.with_suffix('.docx')

        self.doc.save(output_path)
        return str(output_path)


def convert_doc(input_file: str, output_file: str | None = None) -> str:
    """便捷函数：转换文档"""
    converter = DisDocConverter()
    return converter.convert_file(input_file, output_file)


def convert_content(markdown: str, output_file: str) -> str:
    """便捷函数：转换内容"""
    converter = DisDocConverter()
    converter.convert(markdown)
    converter.doc.save(output_file)
    return output_file


def main():
    """命令行入口"""
    sys.stdout.reconfigure(encoding='utf-8')

    if len(sys.argv) < 2:
        print(__doc__)
        print("\nUsage:")
        print("  python doc_converter.py input.md [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        result = convert_doc(input_file, output_file)
        print(f"[OK] Converted: {result}")
    except Exception as e:
        print(f"[FAIL] Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()